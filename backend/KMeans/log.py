from docx import shared
import docx
import datetime
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path


def init():
    my_file = Path("documents/log.docx")
    if my_file.is_file():
        pass
    else:
        doc = docx.Document()
        paragraph = doc.add_heading("程序结果日志")
        paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.save(my_file)


def addpicture(path):
    fn = r'documents/log.docx'
    doc = docx.Document(fn)
    i = datetime.datetime.now()
    doc.add_paragraph(str(i))
    doc.add_picture(path, height=shared.Cm(14))
    doc.save("documents/log.docx")


def aaddtext(data):
    fn = r'documents/log.docx'
    doc = docx.Document(fn)
    i = datetime.datetime.now()
    doc.add_paragraph(str(i))
    doc.add_paragraph(data)
    doc.save("documents/log.docx")


def clear():
    my_file = Path("documents/log.docx")
    doc = docx.Document()
    paragraph = doc.add_heading("程序结果日志")
    paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.save(my_file)